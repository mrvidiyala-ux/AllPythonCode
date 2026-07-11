from docx import Document

# Create the Word document
doc = Document()
doc.add_heading('Italian Brainrot Funny Names', level=1)

names = [
    "Spaghettino Bambolino",
    "Bingus Raviolino",
    "Gabagoolo Supremo",
    "Mozzarellius Maximus",
    "Tony Pepperonissimo",
    "Rigatoni McFettuccine",
    "Luigi Bananaroni",
    "Pizzarino Explosione",
    "Don Waffliano",
    "Fettuccino Braincello",
    "Mamma Mia Giuseppe Jr.",
    "Bolognese Bonkoni",
    "Alfredo Spinnypants",
    "Panini Lamborghini",
    "Vincenzo Spaghetti-Face",
    "Macaronio Delusione",
    "Sir Prosciutto Madness",
    "Cannoli Overdrive",
    "Giovanni Oops-a-Mia",
    "Tortellini Goblinio"
]

for name in names:
    doc.add_paragraph(name, style='List Bullet')

# Save the document
file_path = "Italian_Brainrot_Funny_Names.docx"
doc.save(file_path)

file_path
